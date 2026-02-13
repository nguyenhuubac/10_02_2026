import streamlit as st
import google.generativeai as genai
import pdfplumber
from docx import Document
import requests
import json
import random
import time
from io import BytesIO

# --- C·∫§U H√åNH LI√äN K·∫æT FILE KEY ---
DRIVE_FILE_LINK = "https://drive.google.com/file/d/1iBZqNSs6VyhFB5hQldG_5XBPKFtMfGuV/view?usp=sharing"

# --- C·∫§U H√åNH MODEL THEO Y√äU C·∫¶U ---
# H·ªá th·ªëng s·∫Ω th·ª≠ l·∫ßn l∆∞·ª£t t·ª´ tr√°i sang ph·∫£i
TARGET_MODELS = ["gemini-3-pro-preview", "gemini-3-flash-preview"]

# --- C·∫§U H√åNH TRANG ---
st.set_page_config(page_title="Tr·ª£ l√Ω Gi√°o √°n Gen-3", page_icon="‚ö°", layout="wide")

st.title(f"‚ö° Tr·ª£ l√Ω Th·∫©m ƒë·ªãnh Gi√°o √°n")
#st.markdown(f"ƒêang k√≠ch ho·∫°t ch·∫ø ƒë·ªô Preview cho: **{', '.join(TARGET_MODELS)}**")

# --- H√ÄM X·ª¨ L√ù FILE ---
def read_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def read_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# --- H√ÄM T·∫¢I KEY T·ª™ DRIVE ---
def load_keys_from_drive(url):
    try:
        file_id = url.split('/d/')[1].split('/')[0]
        download_url = f'https://drive.google.com/uc?export=download&id={file_id}'
        
        response = requests.get(download_url)
        if response.status_code == 200:
            keys = json.loads(response.text)
            if isinstance(keys, list) and len(keys) > 0:
                return keys
            else:
                st.error("File JSON kh√¥ng ƒë√∫ng ƒë·ªãnh d·∫°ng danh s√°ch!")
                return []
        else:
            st.error(f"L·ªói t·∫£i file (Code: {response.status_code})")
            return []
    except Exception as e:
        st.error(f"L·ªói x·ª≠ l√Ω link Drive: {e}")
        return []

from io import BytesIO # Th√™m th∆∞ vi·ªán x·ª≠ l√Ω file tr√™n RAM

# --- H√ÄM G·ªåI GEMINI ƒêA T·∫¶NG (MULTI-MODEL RETRY) ---
def generate_content_with_retry(prompt, keys_list):
    available_keys = keys_list.copy()
    random.shuffle(available_keys) 
    
    # 1. V√≤ng l·∫∑p qua t·ª´ng Key
    for api_key in available_keys:
        genai.configure(api_key=api_key)
        
        # 2. V√≤ng l·∫∑p qua t·ª´ng Model (∆Øu ti√™n Pro -> Flash)
        for model_name in TARGET_MODELS:
            try:
                # C·∫•u h√¨nh generation
                generation_config = {
                    "temperature": 0.7,
                    "top_p": 0.95,
                    "top_k": 64,
                    "max_output_tokens": 8192,
                }
                
                model = genai.GenerativeModel(
                    model_name=model_name,
                    generation_config=generation_config
                )
                
                # G·ªçi API
                print(f"ƒêang th·ª≠ Key: ...{api_key[-5:]} v·ªõi Model: {model_name}")
                response = model.generate_content(prompt)
                
                # N·∫øu ch·∫°y ƒë·∫øn ƒë√¢y l√† th√†nh c√¥ng -> Tr·∫£ v·ªÅ k·∫øt qu·∫£ ngay
                return response.text, api_key, model_name
                
            except Exception as e:
                # In l·ªói v√† th·ª≠ model ti·∫øp theo (ho·∫∑c key ti·∫øp theo)
                print(f"L·ªói [Key: ...{api_key[-5:]}] [Model: {model_name}]: {e}")
                time.sleep(0.5) 
                continue 
            
    # N·∫øu ch·∫°y h·∫øt t·∫•t c·∫£ Key v√† Model m√† v·∫´n l·ªói
    raise Exception("T·∫•t c·∫£ API Key v√† Model ƒë·ªÅu th·∫•t b·∫°i! Vui l√≤ng ki·ªÉm tra l·∫°i quy·ªÅn truy c·∫≠p Preview.")

import re  # <--- NH·ªö TH√äM D√íNG N√ÄY L√äN TR√äN C√ôNG FILE C·ª¶A B·∫†N

# --- H√ÄM T·∫†O FILE WORD K·∫æT QU·∫¢ (B·∫£n "S·∫°ch" & X·ª≠ l√Ω To√°n) ---
def create_docx_file(text_content):
    doc = Document()
    doc.add_heading('K·∫æT QU·∫¢ TH·∫®M ƒê·ªäNH GI√ÅO √ÅN', 0)
    
    # Danh s√°ch c√°c t·ª´ Ti·∫øng Anh c·∫ßn lo·∫°i b·ªè
    # B·∫°n c√≥ th·ªÉ th√™m c√°c t·ª´ kh√°c v√†o danh s√°ch n√†y
    english_words_to_remove = ["(Flow)", "(Creative Corner)", "Wow", "Creative Corner"]
    
    is_main_content = False
    
    for line in text_content.split('\n'):
        line = line.strip()
        
        # 1. B·ªè d√≤ng tr·∫Øng v√† d√≤ng g·∫°ch ngang (---)
        if not line or line == '---' or line.startswith('---'):
            continue
            
        # 2. Logic b·ªè qua ƒëo·∫°n ch√†o h·ªèi ƒë·∫ßu (Ch·ªâ l·∫•y t·ª´ m·ª•c 1)
        if line.startswith("## 1"):
            is_main_content = True
        if not is_main_content:
            continue

        # 3. X·ª¨ L√ù L√ÄM S·∫†CH VƒÇN B·∫¢N (QUAN TR·ªåNG)
        
        # a. X√≥a c√°c t·ª´ ti·∫øng Anh trong danh s√°ch
        for word in english_words_to_remove:
            line = line.replace(word, "")
            
        # b. X·ª≠ l√Ω c√¥ng th·ª©c To√°n (LaTeX): Bi·∫øn $m$ th√†nh m
        # D√πng Regex t√¨m t·∫•t c·∫£ nh·ªØng g√¨ n·∫±m trong $...$ v√† b·ªè d·∫•u $ ƒëi
        line = re.sub(r'\$(.*?)\$', r'\1', line)

        # c. X√≥a d·∫•u b√¥i ƒë·∫≠m (** v√† __)
        clean_line = line.replace('**', '').replace('__', '').strip()

        # 4. GHI V√ÄO FILE WORD
        if clean_line.startswith('## '):
            # Ti√™u ƒë·ªÅ 1 (B·ªè d·∫•u ##)
            doc.add_heading(clean_line.replace('## ', ''), level=1)
            
        elif clean_line.startswith('### '):
            # Ti√™u ƒë·ªÅ 2 (B·ªè d·∫•u ###)
            doc.add_heading(clean_line.replace('### ', ''), level=2)
            
        elif clean_line.startswith('* ') or clean_line.startswith('- '):
            # G·∫°ch ƒë·∫ßu d√≤ng: B·ªè d·∫•u * ho·∫∑c - v√† g√°n style Bullet
            content_text = clean_line[2:].strip()
            # N·∫øu sau khi x√≥a d·∫•u * m√† d√≤ng b·ªã r·ªóng th√¨ b·ªè qua
            if content_text:
                doc.add_paragraph(content_text, style='List Bullet')
                
        else:
            # ƒêo·∫°n vƒÉn th∆∞·ªùng
            if clean_line:
                doc.add_paragraph(clean_line)
            
    # L∆∞u file v√†o RAM
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio# --- GIAO DI·ªÜN CH√çNH ---
col1, col2 = st.columns([1, 2]) 

with col1:
    st.info("üìÇ **T·∫£i file**")
    uploaded_file = st.file_uploader("Ch·ªçn gi√°o √°n (PDF/Word)", type=["pdf", "docx"])

if uploaded_file is not None:
    # 1. ƒê·ªçc file
    file_type = uploaded_file.name.split(".")[-1]
    try:
        if file_type == "pdf":
            content = read_pdf(uploaded_file)
        elif file_type == "docx":
            content = read_docx(uploaded_file)
        
        with col1:
            with st.expander("Xem n·ªôi dung th√¥"):
                st.text_area("Tr√≠ch xu·∫•t:", content[:1000] + "...", height=300)

        # 2. N√∫t ki·ªÉm tra
        with col2:
            st.success(f"‚úÖ File ƒë√£ nh·∫≠n! S·∫µn s√†ng th·ª≠ nghi·ªám tr√™n {len(TARGET_MODELS)} models.")
            if st.button("üîç Ki·ªÉm tra", type="primary"):
                with st.status("ƒêang kh·ªüi ƒë·ªông ti·∫øn tr√¨nh x·ª≠ l√Ω...", expanded=True) as status:
                    
                    status.write("üîÑ ƒêang l·∫•y danh s√°ch")
                    keys_list = load_keys_from_drive(DRIVE_FILE_LINK)
                    
                    if keys_list:
                        status.write(f"üîê ƒê√£ n·∫°p {len(keys_list)} Key. ƒêang k·∫øt n·ªëi...")
                        
                        prompt = f"""
                        B·∫°n l√† Chuy√™n gia S∆∞ ph·∫°m (S·ª≠ d·ª•ng model Gemini 3 th·∫ø h·ªá m·ªõi).
                        H√£y ph√¢n t√≠ch gi√°o √°n sau ƒë√¢y. Y√™u c·∫ßu t∆∞ duy logic s√¢u chu·ªói, ph√°t hi·ªán l·ªói ·∫©n v√† g·ª£i √Ω s√°ng t·∫°o.

                        N·ªòI DUNG GI√ÅO √ÅN:
                        {content}

                        Y√äU C·∫¶U OUTPUT (Markdown):
                        
                        ## 1. T·ªïng quan
                        * ƒê√°nh gi√° ch·∫•t l∆∞·ª£ng: .../10
                        * Nh·∫≠n ƒë·ªãnh chung: ...

                        ## 2. Ph√¢n t√≠ch S√¢u
                        * **M·ª•c ti√™u:** Ph√¢n t√≠ch k·ªπ t√≠nh kh·∫£ thi v√† ƒë·ªãnh l∆∞·ª£ng.
                        * **Ho·∫°t ƒë·ªông:** Ph√¢n t√≠ch d√≤ng ch·∫£y t∆∞ duy (Flow) c·ªßa h·ªçc sinh.
                        * **C√¥ng ngh·ªá:** ƒê√°nh gi√° vi·ªác ·ª©ng d·ª•ng CNTT/AI trong b√†i (n·∫øu c√≥).

                        ## 3. C√°c l·ªói c·∫ßn kh·∫Øc ph·ª•c ngay
                        * ...

                        ## 4. G√≥c S√°ng t·∫°o
                        * ƒê·ªÅ xu·∫•t 1 ho·∫°t ƒë·ªông thay th·∫ø "Wow" ƒë·ªÉ g√¢y ·∫•n t∆∞·ª£ng m·∫°nh cho h·ªçc sinh.
                        """
                        
                        try:
                            # G·ªçi h√†m x·ª≠ l√Ω ƒëa t·∫ßng
                            result_text, used_key, used_model = generate_content_with_retry(prompt, keys_list)
                            
                            status.update(label="Ho√†n t·∫•t!", state="complete", expanded=False)
                            st.balloons()
                            
                            # Hi·ªÉn th·ªã k·∫øt qu·∫£
                            st.markdown(f"### K·∫øt qu·∫£ ph√¢n t√≠ch t·ª´: `{used_model}`")
                            st.markdown(result_text)
                            st.caption(f"ƒê√£ x·ª≠ l√Ω th√†nh c√¥ng b·ªüi Key: ...{used_key[-6:]}")
                            
                            # --- XU·∫§T FILE WORD (M·ªõi th√™m) ---
                            st.divider()
                            docx_file = create_docx_file(result_text)
                            st.download_button(
                                label="üì• T·∫£i k·∫øt qu·∫£ v·ªÅ (Word)",
                                data=docx_file,
                                file_name="Ket_qua_tham_dinh_giao_an.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary"
                            )
                            # ---------------------------------
                            
                        except Exception as e:
                            status.update(label="Th·∫•t b·∫°i!", state="error")
                            st.error(f"L·ªói h·ªá th·ªëng: {e}")
                    else:
                        status.update(label="L·ªói Key!", state="error")

    except Exception as e:
        st.error(f"L·ªói ƒë·ªçc file: {e}")

else:
    with col2:
        st.info("üëà M·ªùi b·∫°n t·∫£i gi√°o √°n l√™n")